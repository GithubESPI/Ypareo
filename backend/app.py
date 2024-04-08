from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
from docx import Document
import os

app = FastAPI()

# Configuration de CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Assurez-vous que les dossiers existent
os.makedirs("uploads", exist_ok=True)
os.makedirs("bulletins/M1_S1", exist_ok=True)

@app.post("/upload-excel/")
async def upload_excel(file: UploadFile = File(...)):
    try:
        file_location = f"uploads/{file.filename}"
        with open(file_location, "wb+") as file_object:
            file_object.write(file.file.read())
        
        # Appel de la fonction de traitement du fichier Excel
        bulletin_paths = process_excel_file(file_location)

        # Retour d'une réponse de succès avec les chemins des bulletins générés
        return {"filenames": bulletin_paths}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erreur de traitement du fichier: {e}")

def process_excel_file(file_path: str):
    df = pd.read_excel(file_path, header=4)
    bulletin_paths = []
    for index, row in df.iterrows():
        bulletin_path = generate_word_document(row)
        bulletin_paths.append(bulletin_path)
    return bulletin_paths

def replace_paragraph_text(paragraph, text):
    """Remplace le texte d'un paragraphe tout en conservant la mise en forme."""
    if paragraph.text:
        p_element = paragraph._element
        p_element.clear_content()
        p_element.add_run(text)

def generate_word_document(data):
    template_path = "bulletins/M1_S1/modele_bulletin_M1_S1.docx"
    document = Document(template_path)
    
    for paragraph in document.paragraphs:
        if "{{nom_apprenant}}" in paragraph.text:
            replace_paragraph_text(paragraph, paragraph.text.replace("{{nom_apprenant}}", data.get("Nom", "Inconnu")))
        # Ajoutez ici des conditions supplémentaires pour d'autres placeholders

    bulletin_filename = f"bulletins/M1_S1/{data.get('Nom', 'Inconnu').replace(' ', '_')}_bulletin.docx"
    document.save(bulletin_filename)
    
    return bulletin_filename

@app.get("/download-bulletin/{student_name}")
async def download_bulletin(student_name: str):
    file_path = f"bulletins/M1_S1/{student_name}_bulletin.docx"
    if os.path.exists(file_path):
        return FileResponse(path=file_path, filename=f"{student_name}_bulletin.docx")
    else:
        raise HTTPException(status_code=404, detail="Bulletin non trouvé")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
